
import os, sqlite3, hashlib, json, yaml, base64, datetime as dt
from typing import Dict, List, Tuple, Optional
import pandas as pd
from docx import Document

# cryptography
try:
    from cryptography.hazmat.primitives import hashes, serialization
    from cryptography.hazmat.primitives.asymmetric import padding, rsa
    from cryptography.hazmat.backends import default_backend
    from cryptography import x509
    from cryptography.x509.oid import NameOID
except Exception:
    hashes = serialization = padding = rsa = default_backend = x509 = NameOID = None

# ---------- Hash utils ----------
def sha256_of_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()

def chain_hash(prev_hash: str, payload: Dict) -> str:
    data = (prev_hash or "") + json.dumps(payload, sort_keys=True, ensure_ascii=False)
    return sha256_of_text(data)

# ---------- PKI (RSA) ----------
def generate_rsa_keypair(bits: int = 2048):
    if rsa is None:
        raise RuntimeError("cryptography no disponible")
    private_key = rsa.generate_private_key(public_exponent=65537, key_size=bits, backend=default_backend())
    priv_pem = private_key.private_bytes(
        encoding=serialization.Encoding.PEM,
        format=serialization.PrivateFormat.PKCS8,
        encryption_algorithm=serialization.NoEncryption()
    ).decode("utf-8")
    pub_pem = private_key.public_key().public_bytes(
        encoding=serialization.Encoding.PEM,
        format=serialization.PublicFormat.SubjectPublicKeyInfo
    ).decode("utf-8")
    return priv_pem, pub_pem

def sign_with_pem(private_pem: str, message: str) -> str:
    if serialization is None:
        raise RuntimeError("cryptography no disponible")
    private_key = serialization.load_pem_private_key(private_pem.encode("utf-8"), password=None, backend=default_backend())
    sig = private_key.sign(
        message.encode("utf-8"),
        padding.PSS(mgf=padding.MGF1(hashes.SHA256()), salt_length=padding.PSS.MAX_LENGTH),
        hashes.SHA256()
    )
    return base64.b64encode(sig).decode("utf-8")

def verify_with_pem(public_pem: str, message: str, b64sig: str) -> bool:
    if serialization is None:
        raise RuntimeError("cryptography no disponible")
    public_key = serialization.load_pem_public_key(public_pem.encode("utf-8"), backend=default_backend())
    try:
        public_key.verify(
            base64.b64decode(b64sig.encode("utf-8")),
            message.encode("utf-8"),
            padding.PSS(mgf=padding.MGF1(hashes.SHA256()), salt_length=padding.PSS.MAX_LENGTH),
            hashes.SHA256()
        )
        return True
    except Exception:
        return False

# ---------- TSA demo (RFC3161-like, local) ----------
def tsa_generate_token(private_pem: str, imprint_hex: str, policy: str = "1.2.3.4.5-demo") -> Dict:
    """
    Genera un token de sello de tiempo *demo* (no oficial).
    Estructura: { 'imprint': hex, 'policy': OID string, 'tsa_ts': ISO, 'nonce': hex, 'sig_b64': base64 }
    Firma = RSA-PSS(SHA256) sobre JSON canonizado de {imprint, policy, tsa_ts, nonce}
    """
    payload = {
        "imprint": imprint_hex.lower(),
        "policy": policy,
        "tsa_ts": pd.Timestamp.utcnow().isoformat(),
        "nonce": sha256_of_text(os.urandom(16).hex())[:32]
    }
    msg = json.dumps(payload, sort_keys=True, ensure_ascii=False)
    sig_b64 = sign_with_pem(private_pem, msg)
    payload["sig_b64"] = sig_b64
    return payload

def tsa_verify_token(public_pem: str, token: Dict) -> bool:
    core = {k: token[k] for k in ["imprint","policy","tsa_ts","nonce"]}
    msg = json.dumps(core, sort_keys=True, ensure_ascii=False)
    return verify_with_pem(public_pem, msg, token.get("sig_b64",""))

# ---------- X.509 helpers ----------
def parse_certificate(pem: str) -> Optional[Dict]:
    if x509 is None:
        raise RuntimeError("cryptography no disponible")
    try:
        cert = x509.load_pem_x509_certificate(pem.encode("utf-8"), default_backend())
        subject = cert.subject.get_attributes_for_oid(NameOID.COMMON_NAME)[0].value if cert.subject.get_attributes_for_oid(NameOID.COMMON_NAME) else ""
        issuer = cert.issuer.get_attributes_for_oid(NameOID.COMMON_NAME)[0].value if cert.issuer.get_attributes_for_oid(NameOID.COMMON_NAME) else ""
        not_before = cert.not_valid_before
        not_after = cert.not_valid_after
        sn = cert.serial_number
        return {"subject_cn": subject, "issuer_cn": issuer, "not_before": not_before.isoformat(), "not_after": not_after.isoformat(), "serial": hex(sn), "cert_obj": cert}
    except Exception:
        return None

def verify_cert_signature(issuer_pem: str, cert_pem: str) -> bool:
    """Verifica que issuer haya firmado cert (firma básica, sin OCSP/CRL)."""
    try:
        issuer = x509.load_pem_x509_certificate(issuer_pem.encode("utf-8"), default_backend())
        cert = x509.load_pem_x509_certificate(cert_pem.encode("utf-8"), default_backend())
        pub = issuer.public_key()
        pub.verify(cert.signature, cert.tbs_certificate_bytes, padding.PKCS1v15(), cert.signature_hash_algorithm)
        return True
    except Exception:
        return False

def is_cert_time_valid(cert_pem: str, at_utc_iso: Optional[str] = None) -> bool:
    try:
        cert = x509.load_pem_x509_certificate(cert_pem.encode("utf-8"), default_backend())
        now = pd.Timestamp.utcnow() if at_utc_iso is None else pd.Timestamp(at_utc_iso)
        return cert.not_valid_before <= now <= cert.not_valid_after
    except Exception:
        return False

# ---------- Core econ (igual previo) ----------
def zone_from_P(P: float, T1: float, T2: float) -> str:
    if P <= T1: return "elastica"
    if P <= T2: return "fluencia"
    return "ruptura"

def evaluate_scenarios(df: pd.DataFrame, T1: float, T2: float, E_c: float, E_v: float, weights: Dict[str,float], policies: Dict[str,float], cost_rel: float=0.78, util_params: Dict[str, Dict[str,float]] = None) -> pd.DataFrame:
    E_joint = (E_c + E_v)/2.0
    out = df.copy()
    out["precio_rel"] = out["precio_rel"].astype(float); out["plazo_rel"] = out["plazo_rel"].astype(float)
    out["alcance_gap"] = out["alcance_gap"].astype(float); out["riesgo"] = out["riesgo"].astype(float)
    gpre = (1.0 - out["precio_rel"]).abs(); gpla = (1.0 - out["plazo_rel"]).abs(); galc = out["alcance_gap"].abs()
    out["P"] = (weights["precio"]*gpre + weights["plazo"]*gpla + weights["alcance"]*galc + weights["riesgo"]*out["riesgo"]) / max(E_joint, 1e-6)
    out["M"] = (E_c + E_v) - (gpre + gpla + out["riesgo"]); out["zona"] = out["P"].apply(lambda x: zone_from_P(x, T1, T2))
    value_rel = 1.05 - 0.4*galc; out["ROI_c"] = (value_rel / out["precio_rel"]) - 1.0 - 0.05*(out["plazo_rel"] - 1.0).clip(lower=0); out["ROI_c"] = out["ROI_c"].clip(lower=-1.0)
    out["margen_v"] = ((out["precio_rel"] - cost_rel)/out["precio_rel"]).clip(lower=-1.0)
    bp = util_params.get("buyer", {"alpha":1.0,"beta":0.6,"gamma":0.7}) if util_params else {"alpha":1.0,"beta":0.6,"gamma":0.7}
    sp = util_params.get("seller", {"alpha":1.0,"beta":0.6,"gamma":0.7}) if util_params else {"alpha":1.0,"beta":0.6,"gamma":0.7}
    out["buyer_util"]  = bp["alpha"]*out["ROI_c"]    - bp["beta"]*out["riesgo"] - bp["gamma"]*out["P"]
    out["seller_util"] = sp["alpha"]*out["margen_v"] - sp["beta"]*out["riesgo"] - sp["gamma"]*out["P"]
    out["feasible"] = (out["P"]<=T2) & (out["M"]>0) & (out["ROI_c"]>=policies.get("roi_min_c",0.10)) & (out["margen_v"]>=policies.get("margen_min_v",0.15))
    arr = out[["buyer_util","seller_util","feasible"]].to_numpy(); pareto = []
    for i in range(len(arr)):
        bu_i, su_i, feas_i = arr[i]
        if not feas_i: pareto.append(False); continue
        dominated = False
        for j in range(len(arr)):
            if i==j: continue
            bu_j, su_j, feas_j = arr[j]
            if not feas_j: continue
            if (bu_j >= bu_i and su_j >= su_i) and (bu_j > bu_i or su_j > su_i): dominated = True; break
        pareto.append(not dominated)
    out["pareto"] = pareto; return out

def load_vertical_rules(path: str) -> dict:
    if not os.path.exists(path): return {}
    with open(path, "r", encoding="utf-8") as f: return yaml.safe_load(f) or {}

def recommend_compensations(P: float, M: float, zone: str, gaps: Dict[str,float]) -> List[str]:
    recs = []
    if zone == "elastica":
        recs.append("Proceder y documentar; mantener vigilancia ligera")
    elif zone == "fluencia":
        dom = max(gaps, key=gaps.get)
        if dom == "pre":
            recs.append("Si mantienes precio: reducir alcance o pedir anticipo 20%"); recs.append("Si bajas precio 2-4%: pide SLA L2 y permanencia 12m")
        elif dom == "pla":
            recs.append("Reducir plazo 5-10% o pedir hito intermedio con pago parcial")
        elif dom == "alc":
            recs.append("Ajustar alcance (features menores) o creditos de servicio")
        recs.append("Registrar contrapartidas antes de avanzar")
    else:
        recs.append("Detener y reencuadrar: supera T2 o M<=0"); recs.append("Opciones: dividir proyecto, escalonar pagos, elevar precio o reducir alcance")
    return recs

def recommend_semantic(vertical_rules: dict, vertical: str, gaps: Dict[str,float]) -> List[str]:
    out = []; spec = (vertical_rules.get("verticals") or {}).get(vertical, {}); dmap = (spec.get("dominant_gaps") or {})
    dom = max(gaps, key=gaps.get); out.extend(dmap.get(dom, {}).get("suggestions", []))
    if not out: out.extend(dmap.get("default", {}).get("suggestions", []))
    return out

def route_to_elastic(precio_rel, plazo_rel, alcance_gap, riesgo, T1, E_c, E_v, weights, max_iter=200):
    E_joint = (E_c + E_v)/2.0
    def metrics(pr, pl, ag, r):
        gpre = abs(1.0 - pr); gpla = abs(1.0 - pl)
        P = (weights["precio"]*gpre + weights["plazo"]*gpla + weights["alcance"]*ag + weights["riesgo"]*r) / max(E_joint,1e-6)
        M = (E_c + E_v) - (gpre + gpla + r); return P, M, {"pre":gpre,"pla":gpla,"alc":ag}
    pr, pl, ag, r = precio_rel, plazo_rel, alcance_gap, riesgo
    P, M, _ = metrics(pr, pl, ag, r)
    if P<=T1 and M>0: return (pr, pl, ag, r), {"P":P,"M":M,"delta_sum":0.0,"steps":[]}
    steps_log = []
    for _ in range(max_iter):
        candidates = []
        for var in ["precio_rel","plazo_rel","alcance_gap","riesgo"]:
            npr, npl, nag, nr = pr, pl, ag, r
            if var=="precio_rel": npr = pr - 0.01 if pr>1.0 else pr + 0.01
            elif var=="plazo_rel": npl = pl - 0.01 if pl>1.0 else pl + 0.01
            elif var=="alcance_gap": nag = max(0.0, ag - 0.01)
            elif var=="riesgo": nr = max(0.0, r - 0.01)
            P2, M2, gaps2 = metrics(npr, npl, nag, nr); delta_sum = abs(npr-precio_rel)+abs(npl-plazo_rel)+abs(nag-alcance_gap)+abs(nr-riesgo)
            candidates.append((P2, M2, delta_sum, npr, npl, nag, nr, gaps2, var))
        candidates.sort(key=lambda x: (x[0]>T1 or x[1]<=0, x[0], x[2]))
        Pn, Mn, ds, pr, pl, ag, r, gaps2, moved = candidates[0]
        steps_log.append({"move": moved, "precio_rel": pr, "plazo_rel": pl, "alcance_gap": ag, "riesgo": r, "P":Pn, "M":Mn})
        if Pn<=T1 and Mn>0:
            return (pr, pl, ag, r), {"P":Pn, "M":Mn, "delta_sum":abs(pr-precio_rel)+abs(pl-plazo_rel)+abs(ag-alcance_gap)+abs(r-riesgo), "steps": steps_log}
    return (pr, pl, ag, r), {"P":Pn, "M":Mn, "delta_sum":abs(pr-precio_rel)+abs(pl-plazo_rel)+abs(ag-alcance_gap)+abs(r-riesgo), "steps": steps_log}

# ---------- SQLite (igual v5) ----------
def init_db(db_path: str):
    os.makedirs(os.path.dirname(db_path), exist_ok=True)
    con = sqlite3.connect(db_path); cur = con.cursor()
    cur.execute("""CREATE TABLE IF NOT EXISTS negociacion(id INTEGER PRIMARY KEY AUTOINCREMENT, compra_id TEXT, comprador TEXT, vendedor TEXT, T1 REAL, T2 REAL, E_c REAL, E_v REAL, created_ts TEXT)""")
    cur.execute("""CREATE TABLE IF NOT EXISTS iteracion(id INTEGER PRIMARY KEY AUTOINCREMENT, compra_id TEXT, ts TEXT, precio_rel REAL, plazo_rel REAL, alcance_gap REAL, riesgo REAL, P REAL, M REAL, zona TEXT, acciones_json TEXT)""")
    cur.execute("""CREATE TABLE IF NOT EXISTS cierre(compra_id TEXT PRIMARY KEY, fecha TEXT, resolucion TEXT, contrapartidas_json TEXT, evidencias_json TEXT, sha256 TEXT, buyer_sig_json TEXT, seller_sig_json TEXT)""")
    con.commit(); con.close()

def insert_negociacion(db_path, compra_id, comprador, vendedor, T1, T2, E_c, E_v):
    con = sqlite3.connect(db_path); cur = con.cursor()
    cur.execute("INSERT INTO negociacion(compra_id,comprador,vendedor,T1,T2,E_c,E_v,created_ts) VALUES (?,?,?,?,?,?,?,datetime('now'))",(compra_id,comprador,vendedor,T1,T2,E_c,E_v))
    con.commit(); con.close()

def insert_iteracion(db_path, compra_id, it):
    con = sqlite3.connect(db_path); cur = con.cursor()
    cur.execute("""INSERT INTO iteracion(compra_id,ts,precio_rel,plazo_rel,alcance_gap,riesgo,P,M,zona,acciones_json) VALUES (?,?,?,?,?,?,?,?,?,?)""",
                (compra_id, it.get("t"), it.get("precio_rel"), it.get("plazo_rel"), it.get("alcance_gap"), it.get("riesgo"), it.get("P"), it.get("M"), it.get("zona"), json.dumps(it.get("acciones", []), ensure_ascii=False)))
    con.commit(); con.close()

def insert_cierre(db_path, compra_id, fecha, resolucion, contrapartidas, evidencias, sha256, buyer_sig=None, seller_sig=None):
    con = sqlite3.connect(db_path); cur = con.cursor()
    cur.execute("""INSERT OR REPLACE INTO cierre(compra_id,fecha,resolucion,contrapartidas_json,evidencias_json,sha256,buyer_sig_json,seller_sig_json) VALUES (?,?,?,?,?,?,?,?)""",
                (compra_id, fecha, resolucion, json.dumps(contrapartidas, ensure_ascii=False), json.dumps(evidencias, ensure_ascii=False), sha256, json.dumps(buyer_sig or {}, ensure_ascii=False), json.dumps(seller_sig or {}, ensure_ascii=False)))
    con.commit(); con.close()

# ---------- Acta / Contrato DOCX ----------
def make_acta_cierre(context: Dict, iter_log, cierre: Dict):
    header = "# Acta de Cierre - Compra Viva AUP\n\nFecha: " + cierre.get("fecha","") + "\n\n"
    meta = "**Compra ID:** " + str(context.get("compra_id","")) + "\n" + "**Comprador:** " + str(context.get("comprador","")) + "\n" + "**Vendedor:** " + str(context.get("vendedor","")) + "\n" + "**Umbrales:** T1=" + str(context.get("T1")) + ", T2=" + str(context.get("T2")) + "\n" + "**Elasticidades:** E_c=" + str(context.get("E_c")) + ", E_v=" + str(context.get("E_v")) + "\n"
    body = "## Iteraciones registradas\n\n"
    for it in iter_log:
        body += "- t: " + str(it.get("t")) + " - precio_rel: " + str(it.get("precio_rel")) + ", plazo_rel: " + str(it.get("plazo_rel")) + ", alcance_gap: " + str(it.get("alcance_gap")) + ", riesgo: " + str(it.get("riesgo")) + ", P: " + str(round(it.get("P",0),3)) + ", M: " + str(round(it.get("M",0),3)) + ", zona: " + str(it.get("zona")) + "\n  - acciones: " + ", ".join(it.get("acciones", [])) + "\n"
    closing = "\n## Cierre\n\n" + "**Resolucion:** " + str(cierre.get("resolucion","")) + "\n" + "**Contrapartidas:** " + ", ".join(cierre.get("contrapartidas", [])) + "\n" + "**Evidencias:** " + ", ".join(cierre.get("evidencias", [])) + "\n"
    md = header + meta + "\n" + body + "\n" + closing
    return md, sha256_of_text(md)

def export_acta_docx(path: str, context: Dict, iter_log, cierre: Dict, sha256: str, buyer_sig=None, seller_sig=None, tsa_token: Dict=None):
    doc = Document(); doc.add_heading('Acta de Cierre - Compra Viva AUP', level=1)
    doc.add_paragraph(f"Fecha: {cierre.get('fecha','')}"); doc.add_paragraph(f"Compra ID: {context.get('compra_id','')}"); doc.add_paragraph(f"Comprador: {context.get('comprador','')}"); doc.add_paragraph(f"Vendedor: {context.get('vendedor','')}")
    doc.add_paragraph(f"Umbrales: T1={context.get('T1')} T2={context.get('T2')}  Elasticidades: E_c={context.get('E_c')} E_v={context.get('E_v')}")
    doc.add_heading('Iteraciones', level=2); table = doc.add_table(rows=1, cols=8); hdr = table.rows[0].cells
    hdr[0].text='t'; hdr[1].text='precio_rel'; hdr[2].text='plazo_rel'; hdr[3].text='alcance_gap'; hdr[4].text='riesgo'; hdr[5].text='P'; hdr[6].text='M'; hdr[7].text='zona'
    for it in iter_log:
        row = table.add_row().cells
        row[0].text=str(it.get('t')); row[1].text=f"{it.get('precio_rel')}"; row[2].text=f"{it.get('plazo_rel')}"; row[3].text=f"{it.get('alcance_gap')}"; row[4].text=f"{it.get('riesgo')}"; row[5].text=f"{round(it.get('P',0),3)}"; row[6].text=f"{round(it.get('M',0),3)}"; row[7].text=str(it.get('zona'))
    doc.add_heading('Cierre', level=2); doc.add_paragraph(f"Resolucion: {cierre.get('resolucion','')}"); doc.add_paragraph('Contrapartidas: ' + ', '.join(cierre.get('contrapartidas', []))); doc.add_paragraph('Evidencias: ' + ', '.join(cierre.get('evidencias', []))); doc.add_paragraph(f"SHA256: {sha256}")
    if buyer_sig: doc.add_paragraph(f"Firma comprador (PKI/Hash): {json.dumps(buyer_sig)[:300]}...")
    if seller_sig: doc.add_paragraph(f"Firma vendedor (PKI/Hash): {json.dumps(seller_sig)[:300]}...")
    if tsa_token: doc.add_paragraph("TSA demo: " + json.dumps(tsa_token)[:400] + "...")
    doc.save(path)

def export_contract_docx(path: str, context: Dict, cierre: Dict, terms: Dict, buyer_sig=None, seller_sig=None):
    doc = Document(); doc.add_heading('Contrato de Prestación de Servicios — Modelo AUP', level=1)
    doc.add_paragraph(f"Fecha: {cierre.get('fecha','')}")
    doc.add_paragraph(f"Partes: {context.get('comprador','')} (Comprador) y {context.get('vendedor','')} (Proveedor).")
    doc.add_heading('Objeto y Alcance', level=2); doc.add_paragraph(terms.get('objeto', 'Prestación de servicios bajo parámetros AUP.'))
    doc.add_paragraph(f"Alcance resumido: {terms.get('alcance','(definir)')}")
    doc.add_heading('Condiciones Comerciales', level=2)
    doc.add_paragraph(f"Precio total: {terms.get('precio_total','(definir)')} {terms.get('moneda','MXN')}  |  Anticipo: {terms.get('anticipo_pct','0')}%")
    doc.add_paragraph(f"Vigencia: {terms.get('vigencia_meses','12')} meses  |  SLA: {terms.get('SLA','L1/L2')}  | Penalizaciones: {terms.get('penalizaciones','(definir)')}")
    doc.add_heading('Gobernanza AUP', level=2)
    doc.add_paragraph(f"Umbrales: T1={context.get('T1')}  T2={context.get('T2')}  |  Elasticidades: E_c={context.get('E_c')}  E_v={context.get('E_v')}")
    doc.add_paragraph("Las partes reconocen el uso de la bitácora viva con hash encadenado por iteración.")
    doc.add_heading('Firmas', level=2)
    doc.add_paragraph(f"Comprador: {buyer_sig or '__________________'}")
    doc.add_paragraph(f"Proveedor: {seller_sig or '__________________'}")
    doc.save(path)
